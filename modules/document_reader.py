"""
LADA v10.0 - Document Reader Module
Read and analyze PDF, Word, and text documents

Features:
- PDF text extraction with PyMuPDF
- Table extraction from PDFs
- Word document (.docx) reading
- AI-powered document summarization
- Search within documents
- Page-by-page extraction
- Image extraction from PDFs
"""

import os
import logging
from pathlib import Path
from typing import Dict, List, Any, Optional, Union
from dataclasses import dataclass
from datetime import datetime

logger = logging.getLogger(__name__)

# Try to import PyMuPDF (fitz)
try:
    import fitz  # PyMuPDF
    PYMUPDF_OK = True
except ImportError:
    fitz = None
    PYMUPDF_OK = False
    logger.warning("PyMuPDF not installed - PDF reading limited. Install with: pip install pymupdf")

# Try to import python-docx for Word documents
try:
    from docx import Document as DocxDocument
    DOCX_OK = True
except ImportError:
    DocxDocument = None
    DOCX_OK = False
    logger.warning("python-docx not installed - DOCX reading disabled")

# Try to import pdfplumber for table extraction
try:
    import pdfplumber
    PDFPLUMBER_OK = True
except ImportError:
    pdfplumber = None
    PDFPLUMBER_OK = False


@dataclass
class DocumentPage:
    """Single page from a document"""
    page_number: int
    text: str
    images: List[Dict] = None
    tables: List[List[List[str]]] = None
    
    def __post_init__(self):
        if self.images is None:
            self.images = []
        if self.tables is None:
            self.tables = []


@dataclass
class DocumentInfo:
    """Document metadata and info"""
    path: str
    title: str
    author: Optional[str]
    page_count: int
    created_date: Optional[str]
    modified_date: Optional[str]
    file_size: int
    format: str  # pdf, docx, txt


@dataclass
class DocumentResult:
    """Result from document reading"""
    success: bool
    info: Optional[DocumentInfo]
    pages: List[DocumentPage]
    full_text: str
    summary: Optional[str] = None
    error: Optional[str] = None


class DocumentReader:
    """
    Universal document reader for LADA.
    Supports PDF, DOCX, and text files.
    """
    
    SUPPORTED_FORMATS = {
        '.pdf': 'pdf',
        '.docx': 'docx',
        '.doc': 'doc',
        '.txt': 'txt',
        '.md': 'markdown',
        '.rtf': 'rtf',
    }
    
    def __init__(self, ai_router=None):
        """
        Initialize document reader.
        
        Args:
            ai_router: Optional HybridAIRouter for AI-powered analysis
        """
        self.ai_router = ai_router
        self._cache = {}  # Cache for recently read documents
        
    def can_read(self, file_path: str) -> bool:
        """Check if file format is supported"""
        ext = Path(file_path).suffix.lower()
        return ext in self.SUPPORTED_FORMATS
    
    def read_document(
        self,
        file_path: str,
        pages: Optional[List[int]] = None,
        extract_tables: bool = True,
        extract_images: bool = False,
        summarize: bool = False
    ) -> DocumentResult:
        """
        Read a document and extract its contents.
        
        Args:
            file_path: Path to the document
            pages: Optional list of page numbers to extract (1-indexed)
            extract_tables: Whether to extract tables
            extract_images: Whether to extract images
            summarize: Whether to generate AI summary
            
        Returns:
            DocumentResult with extracted content
        """
        path = Path(file_path)
        
        if not path.exists():
            return DocumentResult(
                success=False,
                info=None,
                pages=[],
                full_text="",
                error=f"File not found: {file_path}"
            )
        
        ext = path.suffix.lower()
        
        if ext not in self.SUPPORTED_FORMATS:
            return DocumentResult(
                success=False,
                info=None,
                pages=[],
                full_text="",
                error=f"Unsupported format: {ext}"
            )
        
        try:
            if ext == '.pdf':
                result = self._read_pdf(path, pages, extract_tables, extract_images)
            elif ext in ['.docx', '.doc']:
                result = self._read_docx(path)
            elif ext in ['.txt', '.md']:
                result = self._read_text(path)
            else:
                result = self._read_text(path)  # Fallback to text
            
            # Generate summary if requested
            if summarize and result.success and result.full_text:
                result.summary = self._generate_summary(result.full_text)
            
            return result
            
        except Exception as e:
            logger.error(f"Error reading document: {e}")
            return DocumentResult(
                success=False,
                info=None,
                pages=[],
                full_text="",
                error=str(e)
            )
    
    def _read_pdf(
        self,
        path: Path,
        pages: Optional[List[int]],
        extract_tables: bool,
        extract_images: bool
    ) -> DocumentResult:
        """Read PDF document using PyMuPDF"""
        if not PYMUPDF_OK:
            return DocumentResult(
                success=False,
                info=None,
                pages=[],
                full_text="",
                error="PyMuPDF not installed. Run: pip install pymupdf"
            )
        
        doc = fitz.open(str(path))
        
        # Get document info
        metadata = doc.metadata
        info = DocumentInfo(
            path=str(path),
            title=metadata.get('title', path.stem) or path.stem,
            author=metadata.get('author'),
            page_count=len(doc),
            created_date=metadata.get('creationDate'),
            modified_date=metadata.get('modDate'),
            file_size=path.stat().st_size,
            format='pdf'
        )
        
        # Determine which pages to extract
        if pages:
            page_indices = [p - 1 for p in pages if 0 < p <= len(doc)]
        else:
            page_indices = range(len(doc))
        
        extracted_pages = []
        all_text = []
        
        for page_idx in page_indices:
            page = doc[page_idx]
            
            # Extract text
            text = page.get_text("text")
            all_text.append(text)
            
            # Extract images if requested
            images = []
            if extract_images:
                image_list = page.get_images()
                for img_idx, img in enumerate(image_list):
                    images.append({
                        'index': img_idx,
                        'xref': img[0],
                        'width': img[2],
                        'height': img[3]
                    })
            
            # Extract tables if requested
            tables = []
            if extract_tables and PDFPLUMBER_OK:
                try:
                    with pdfplumber.open(str(path)) as pdf:
                        if page_idx < len(pdf.pages):
                            plumber_page = pdf.pages[page_idx]
                            extracted_tables = plumber_page.extract_tables()
                            tables = extracted_tables if extracted_tables else []
                except:
                    pass
            
            extracted_pages.append(DocumentPage(
                page_number=page_idx + 1,
                text=text,
                images=images,
                tables=tables
            ))
        
        doc.close()
        
        return DocumentResult(
            success=True,
            info=info,
            pages=extracted_pages,
            full_text="\n\n".join(all_text)
        )
    
    def _read_docx(self, path: Path) -> DocumentResult:
        """Read Word document using python-docx"""
        if not DOCX_OK:
            return DocumentResult(
                success=False,
                info=None,
                pages=[],
                full_text="",
                error="python-docx not installed. Run: pip install python-docx"
            )
        
        doc = DocxDocument(str(path))
        
        # Get document info
        core_props = doc.core_properties
        info = DocumentInfo(
            path=str(path),
            title=core_props.title or path.stem,
            author=core_props.author,
            page_count=1,  # Word docs don't have fixed pages
            created_date=str(core_props.created) if core_props.created else None,
            modified_date=str(core_props.modified) if core_props.modified else None,
            file_size=path.stat().st_size,
            format='docx'
        )
        
        # Extract all paragraphs
        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text)
        
        # Extract tables
        tables = []
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text for cell in row.cells]
                table_data.append(row_data)
            tables.append(table_data)
        
        full_text = "\n\n".join(paragraphs)
        
        # Create single page (Word docs are treated as one page)
        page = DocumentPage(
            page_number=1,
            text=full_text,
            tables=tables
        )
        
        return DocumentResult(
            success=True,
            info=info,
            pages=[page],
            full_text=full_text
        )
    
    def _read_text(self, path: Path) -> DocumentResult:
        """Read plain text file"""
        # Try different encodings
        encodings = ['utf-8', 'utf-8-sig', 'latin-1', 'cp1252']
        text = None
        
        for encoding in encodings:
            try:
                text = path.read_text(encoding=encoding)
                break
            except UnicodeDecodeError:
                continue
        
        if text is None:
            return DocumentResult(
                success=False,
                info=None,
                pages=[],
                full_text="",
                error="Could not decode file with any supported encoding"
            )
        
        info = DocumentInfo(
            path=str(path),
            title=path.stem,
            author=None,
            page_count=1,
            created_date=None,
            modified_date=datetime.fromtimestamp(path.stat().st_mtime).isoformat(),
            file_size=path.stat().st_size,
            format=path.suffix.lstrip('.') or 'txt'
        )
        
        page = DocumentPage(
            page_number=1,
            text=text
        )
        
        return DocumentResult(
            success=True,
            info=info,
            pages=[page],
            full_text=text
        )
    
    def _generate_summary(self, text: str, max_length: int = 500) -> Optional[str]:
        """Generate AI summary of document text"""
        if not self.ai_router:
            return None
        
        try:
            # Truncate text if too long
            if len(text) > 10000:
                text = text[:10000] + "... [truncated]"
            
            prompt = f"""Summarize the following document in a concise paragraph (max {max_length} characters):

{text}

Summary:"""
            
            response = self.ai_router.query(prompt)
            if response:
                return response[:max_length]
            return None
            
        except Exception as e:
            logger.warning(f"Could not generate summary: {e}")
            return None
    
    def search_document(
        self,
        file_path: str,
        query: str,
        case_sensitive: bool = False
    ) -> List[Dict[str, Any]]:
        """
        Search for text within a document.
        
        Args:
            file_path: Path to the document
            query: Search query
            case_sensitive: Whether search is case sensitive
            
        Returns:
            List of matches with page number and context
        """
        result = self.read_document(file_path)
        
        if not result.success:
            return []
        
        matches = []
        search_query = query if case_sensitive else query.lower()
        
        for page in result.pages:
            text = page.text if case_sensitive else page.text.lower()
            
            # Find all occurrences
            start = 0
            while True:
                pos = text.find(search_query, start)
                if pos == -1:
                    break
                
                # Get context (50 chars before and after)
                context_start = max(0, pos - 50)
                context_end = min(len(page.text), pos + len(query) + 50)
                context = page.text[context_start:context_end]
                
                matches.append({
                    'page': page.page_number,
                    'position': pos,
                    'context': f"...{context}...",
                    'match': page.text[pos:pos + len(query)]
                })
                
                start = pos + 1
        
        return matches
    
    def extract_text_from_pages(
        self,
        file_path: str,
        start_page: int = 1,
        end_page: Optional[int] = None
    ) -> str:
        """
        Extract text from specific page range.
        
        Args:
            file_path: Path to the document
            start_page: Starting page (1-indexed)
            end_page: Ending page (inclusive), None for all remaining
            
        Returns:
            Extracted text
        """
        result = self.read_document(file_path)
        
        if not result.success:
            return ""
        
        if end_page is None:
            end_page = len(result.pages)
        
        pages = [
            p for p in result.pages
            if start_page <= p.page_number <= end_page
        ]
        
        return "\n\n".join(p.text for p in pages)
    
    def get_document_stats(self, file_path: str) -> Dict[str, Any]:
        """
        Get statistics about a document.
        
        Args:
            file_path: Path to the document
            
        Returns:
            Dictionary with document statistics
        """
        result = self.read_document(file_path)
        
        if not result.success:
            return {'error': result.error}
        
        text = result.full_text
        
        # Calculate statistics
        words = text.split()
        sentences = [s for s in text.replace('!', '.').replace('?', '.').split('.') if s.strip()]
        paragraphs = [p for p in text.split('\n\n') if p.strip()]
        
        return {
            'file_name': result.info.title,
            'format': result.info.format,
            'page_count': result.info.page_count,
            'file_size_kb': round(result.info.file_size / 1024, 2),
            'character_count': len(text),
            'word_count': len(words),
            'sentence_count': len(sentences),
            'paragraph_count': len(paragraphs),
            'avg_words_per_sentence': round(len(words) / max(len(sentences), 1), 1),
            'table_count': sum(len(p.tables) for p in result.pages),
            'image_count': sum(len(p.images) for p in result.pages),
        }


# ============================================================
# EXAMPLE USAGE
# ============================================================
if __name__ == "__main__":
    logging.basicConfig(level=logging.INFO)
    
    reader = DocumentReader()
    
    print("📄 Document Reader Test")
    print("=" * 50)
    
    # Test with a sample PDF if exists
    test_files = [
        "test.pdf",
        "document.pdf",
        "README.md"
    ]
    
    for test_file in test_files:
        if Path(test_file).exists():
            print(f"\n📖 Reading: {test_file}")
            result = reader.read_document(test_file)
            
            if result.success:
                print(f"   Title: {result.info.title}")
                print(f"   Pages: {result.info.page_count}")
                print(f"   Text preview: {result.full_text[:200]}...")
                
                # Get stats
                stats = reader.get_document_stats(test_file)
                print(f"   Words: {stats.get('word_count', 0)}")
            else:
                print(f"   Error: {result.error}")
    
    print("\n✅ Document Reader test complete!")
    print("\nSupported formats:", list(DocumentReader.SUPPORTED_FORMATS.keys()))
    print(f"PyMuPDF available: {PYMUPDF_OK}")
    print(f"python-docx available: {DOCX_OK}")
